#!/usr/bin/env python3
"""Generate client-facing API Budget & Requirements DOCX."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "Hybrid-Quant-Engine-API-Budget-Estimate.docx")


def set_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.rows[i + 1].cells[j].text = val
    doc.add_paragraph()


def build() -> Document:
    doc = Document()
    set_defaults(doc)

    title = doc.add_heading("Hybrid Quant Engine", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("External API Budget Estimate & Requirements")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    doc.add_paragraph()

    add_table(doc, ["Field", "Value"], [
        ["Document Version", "1.0"],
        ["Date", "13 July 2026"],
        ["Prepared For", "Client Approval — API & Data Vendor Budget"],
        ["Project", "Hybrid Quant Engine (Analysis, Backtest & Signals)"],
        ["Markets", "India (NSE/BSE) + US (NYSE/NASDAQ)"],
        ["Scope", "Research & signals only — no buy/sell or broker integration"],
        ["Currency", "INR (₹); USD converted at ₹84 / $1 where applicable"],
        ["GST", "18% applicable on Indian vendor invoices (shown separately)"],
    ])

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This document lists all external APIs required to build the Hybrid Quant Engine, "
        "including each vendor name, purpose, technical requirements, and estimated cost. "
        "Costs are presented in two scenarios: MVP (development / pilot) and Production "
        "(client-facing platform). Alert/notification APIs and optional infrastructure "
        "APIs are excluded from this budget."
    )
    add_table(doc, ["Scenario", "Monthly (₹)", "One-Time Year 1 (₹)", "Year 1 Total incl. GST (₹)"], [
        ["MVP — Low estimate", "₹45,872", "₹50,000", "₹7.1 Lakh"],
        ["MVP — High estimate", "₹83,836", "₹1,50,000", "₹13.6 Lakh"],
        ["Production — Low estimate", "₹95,916", "₹1,50,000", "₹15.3 Lakh"],
        ["Production — High estimate", "₹2,25,000", "₹4,00,000", "₹36.6 Lakh"],
    ])
    doc.add_paragraph(
        "Recommendation for client approval: Begin with MVP Low (~₹7.1 Lakh Year 1 incl. GST) "
        "using a 50-symbol pilot universe. Scale to Production after backtest validation."
    )

    doc.add_heading("2. Project Assumptions", level=1)
    assumptions = [
        "10-Point Confluence Scoring Matrix (fundamental + technical + order-flow).",
        "20-year historical backtest window (2006–2026) with 0.25% simulated friction.",
        "Live market data during India shift (IST 09:15–15:30) and US shift (IST 19:00–02:30).",
        "Portal publishes signals/alerts only — no order execution.",
        "Backend: Laravel; AI: Gemini + RAG; DB: MySQL + MongoDB.",
        "TrueData API pricing is estimated (custom quote required from vendor).",
        "Polygon.io and FMP prices per published plans (July 2026).",
    ]
    for a in assumptions:
        doc.add_paragraph(a, style="List Bullet")

    doc.add_heading("3. Detailed API Breakdown by Category", level=1)

    # --- CATEGORY A ---
    doc.add_heading("Category A — India Market Data", level=2)
    doc.add_paragraph("Purpose: Live streaming, EOD history, delivery %, corporate-action-adjusted OHLCV for NSE/BSE.")
    add_table(doc, ["#", "API / Vendor", "Purpose", "Requirements", "MVP Monthly (₹)", "One-Time (₹)"], [
        ["A1", "TrueData Market Data API\n(WebSocket + REST)", "Live India shift streaming; real-time 5min bars; REST historical EOD; DelVolume; bhavcopy", "NSE EQ + BSE EQ; ~50 (MVP) or ~200 (Prod) symbols; WebSocket port 8082; 5min bar stream enabled; REST auth via auth.truedata.in / history.truedata.in; corporate-action-adjusted EOD; delivery data for P9", "8,000 – 20,000", "—"],
        ["A2", "TrueData getBhavCopy (EQ / BSEEQ)", "Bulk daily EOD + delivery post-market ingest", "Segment EQ + BSEEQ; daily automated pull after 18:45 IST; store in MongoDB", "Included in A1", "—"],
        ["A3", "TrueData getBarData (EOD / DelVolume)", "20-year backfill and ongoing daily bars", "Interval EOD + DelVolume; date range 2006–2026; split/bonus adjusted series confirmed with vendor", "Included in A1", "50,000 – 1,50,000\n(MVP backfill)\n1,50,000 – 4,00,000\n(Prod backfill)"],
        ["A4", "NSE / BSE Exchange Data Fees", "Mandatory exchange licensing on top of TrueData", "Non-professional or commercial license per use case; separate from TrueData subscription", "2,000 – 5,000", "—"],
        ["A5", "Exchange Redistribution License\n(Production only)", "Display live/delayed prices on client-facing dashboard", "Written NSE/BSE approval + TrueData redistribution agreement; required if end-users see live data", "0 (MVP internal)\n15,000 – 50,000\n(Production)", "—"],
    ])
    doc.add_paragraph("Category A MVP subtotal: ₹10,000 – ₹25,000/month + ₹50,000 – ₹1,50,000 one-time backfill.")
    doc.add_paragraph("Vendor: https://www.truedata.in/market-data-apis | Quote: sales@truedata.in")

    # --- CATEGORY B ---
    doc.add_heading("Category B — US Market Data", level=2)
    doc.add_paragraph("Purpose: Live US shift streaming, 20Y EOD history, block/large trade data for order-flow scoring (P9).")
    add_table(doc, ["#", "API / Vendor", "Purpose", "Requirements", "MVP Monthly (₹)", "One-Time (₹)"], [
        ["B1", "Polygon.io Stocks Advanced Plan", "Real-time US WebSocket (Shift 2); 20+ years daily OHLCV; trades for block-trade analysis", "NYSE + NASDAQ equities; WebSocket wss://socket.polygon.io; REST api.polygon.io; unlimited API calls; real-time (non-professional tier); ~50–200 symbol universe", "$199 / ₹16,716", "—"],
        ["B2", "Polygon REST Aggregates / Trades", "Historical and live bar construction; volume breakout (P9)", "Daily + minute aggregates; trades endpoint for large/block trade ratio; JSON format", "Included in B1", "—"],
        ["B3", "Polygon Corporate Actions", "Split/dividend adjustment for US backtest", "Historical corp actions per symbol; point-in-time correctness for 20Y walk-forward", "Included in B1", "—"],
        ["B4", "US Exchange Pro Fees\n(if applicable)", "Required if user classified as professional by exchange", "Polygon Non-Pro vs Pro classification; additional fees may apply for commercial redistribution", "0 – 8,400\n(Production)", "—"],
    ])
    doc.add_paragraph("Category B subtotal: ₹16,716 – ₹25,116/month.")
    doc.add_paragraph("Vendor: https://polygon.io/pricing")

    # --- CATEGORY C ---
    doc.add_heading("Category C — Index & Benchmark Data", level=2)
    doc.add_paragraph("Purpose: Mansfield Relative Strength (P7) vs Nifty 50 and S&P 500.")
    add_table(doc, ["#", "API / Vendor", "Purpose", "Requirements", "Monthly (₹)", "One-Time (₹)"], [
        ["C1", "TrueData — NIFTY 50", "India benchmark index history + live", "Symbol NIFTY 50; 20Y EOD + live stream during India shift; same adjustment standard as equities", "Included in A1", "—"],
        ["C2", "Polygon.io — S&P 500 index", "US benchmark index history + live", "Index ticker via Polygon; 20Y daily; live during US shift", "Included in B1", "—"],
    ])
    doc.add_paragraph("Category C subtotal: ₹0 (bundled in Categories A and B).")

    # --- CATEGORY D ---
    doc.add_heading("Category D — Fundamental & Financial Statement Data", level=2)
    doc.add_paragraph("Purpose: Scoring points P1–P5 (ROE, ROIC, Altman Z, Beneish M, Sloan, DSO, CAGR, PEG, FCFF/FCFE).")
    add_table(doc, ["#", "API / Vendor", "Purpose", "Requirements", "MVP Monthly (₹)", "One-Time (₹)"], [
        ["D1", "TrueData Corporate REST API", "India fundamentals, financial statements, ratios", "corporate.getPnLById, getBalSheetById, getCashFlowDetailById; getAllResultsByCompany (quarterly/yearly); standalone + consolidated; quarterly refresh post-results", "8,000 – 15,000", "—"],
        ["D2", "TrueData Corporate WebSocket\n(optional)", "Live corporate announcements during market hours", "rtCorpConnect port 9092; announcement feed for RAG + alert context", "2,000 – 5,000\n(optional)", "—"],
        ["D3", "Financial Modeling Prep (FMP) Premium", "US fundamentals — income statement, balance sheet, cash flow, ratios", "Premium plan: 750 calls/min; 30 years history; full fundamentals; US coverage; commercial use terms reviewed", "$59 / ₹4,956", "—"],
        ["D4", "FMP Ultimate\n(Production upgrade)", "US transcripts, 13F, bulk delivery, higher rate limits", "3,000 calls/min; bulk endpoints; needed at 500+ symbol scale", "$149 / ₹12,516\n(Production)", "—"],
        ["D5", "SimFin Pro API", "US point-in-time fundamentals for honest 20Y backtest", "20+ years history; 30,000 credits/month; avoids lookahead bias in P1–P5 backtest; annual billing $71/mo equivalent", "$71 / ₹5,964\n(recommended)", "—"],
        ["D6", "TrueData getCorpAction", "India corporate actions for price adjustment", "historical.getCorpAction per symbol; splits, bonuses, dividends, rights", "Included in A1/D1", "—"],
    ])
    doc.add_paragraph("Category D MVP subtotal: ₹12,956 – ₹25,920/month (without optional D2).")
    doc.add_paragraph("Vendors: TrueData Corporate | https://site.financialmodelingprep.com/pricing-plans | https://www.simfin.com/en/prices/")

    # --- CATEGORY E ---
    doc.add_heading("Category E — India Compliance & Shareholding Data", level=2)
    doc.add_paragraph("Purpose: Promoter pledge gate (P4) — pass if pledge < 10%, reject if > 15%.")
    add_table(doc, ["#", "API / Vendor", "Purpose", "Requirements", "MVP Monthly (₹)", "One-Time (₹)"], [
        ["E1", "TrueData getSHPListByCompany / getShpDetailById", "Promoter holding and pledge %", "Quarterly shareholding pattern; historical SHP for backtest; symbol-level pledge field", "Included in D1", "—"],
        ["E2", "Trendlyne API\n(fallback only)", "Independent pledge verification if TrueData SHP incomplete", "Current + historical pledge %; quarterly updates; NSE/BSE coverage", "0 – 5,000\n(if needed)", "—"],
    ])
    doc.add_paragraph("Category E subtotal: ₹0 – ₹5,000/month.")

    # --- CATEGORY F ---
    doc.add_heading("Category F — News & Document Ingestion (RAG)", level=2)
    doc.add_paragraph("Purpose: Feed the RAG pipeline for AI research panel — filings, news, announcements (not used in scoring).")
    add_table(doc, ["#", "API / Vendor", "Purpose", "Requirements", "MVP Monthly (₹)", "One-Time (₹)"], [
        ["F1", "SEC EDGAR API", "US 10-K, 10-Q, 8-K filing download", "Free public API; CIK-to-ticker mapping; daily incremental ingest; PDF/HTML storage for RAG chunks", "Free", "—"],
        ["F2", "TrueData getAnnouncementFile / getAnnouncementsList", "India corporate filing PDFs for RAG", "Symbol-filtered announcements; file download by ID; daily ingest job", "Included in D1", "—"],
        ["F3", "Finnhub All-In-One", "Company news feed (India + US)", "300 calls/min; symbol-tagged news; daily RAG refresh; $50/month plan", "$50 / ₹4,200", "—"],
        ["F4", "FMP Market News API", "US press releases and stock news", "Included in FMP Premium/Ultimate subscription; symbol-filtered headlines", "Included in D3", "—"],
        ["F5", "NewsAPI.org Business\n(Production only)", "Broader news coverage for RAG at scale", "250,000 requests/month; commercial license; 5-year article search; CORS enabled", "$449 / ₹37,716\n(Production)", "—"],
    ])
    doc.add_paragraph("Category F MVP subtotal: ₹4,200/month. Production: ₹4,200 – ₹41,916/month.")

    # --- CATEGORY G ---
    doc.add_heading("Category G — AI / LLM / RAG (Gemini)", level=2)
    doc.add_paragraph("Purpose: Research Q&A, filing summaries, score explanations — strictly NOT used for scoring or signals.")
    add_table(doc, ["#", "API / Vendor", "Purpose", "Requirements", "MVP Monthly (₹)", "One-Time (₹)"], [
        ["G1", "Google Gemini 2.5 Flash API", "Chat completion — “Why did X score 7/10?”", "Paid tier; ~50–200 queries/day; $0.30/1M input + $2.50/1M output tokens; no grounding in scoring path", "1,500 – 5,000", "—"],
        ["G2", "Google Gemini Embedding API\n(gemini-embedding-001)", "RAG vector embeddings for filing chunks", "$0.15/1M tokens; batch API $0.075/1M; store vectors in MongoDB; initial corpus embed + daily delta", "500 – 2,000", "—"],
        ["G3", "MongoDB Atlas Vector Search", "Vector storage & similarity search", "Self-hosted MongoDB — no external API fee if on existing cluster; Atlas Vector optional add-on", "0\n(self-hosted)", "—"],
    ])
    doc.add_paragraph("Category G subtotal: ₹2,000 – ₹7,000/month.")
    doc.add_paragraph("Vendor: https://ai.google.dev/gemini-api/docs/pricing")

    doc.add_heading("4. Scoring Point → API Mapping", level=1)
    add_table(doc, ["Score Point", "Description", "Primary API Source"], [
        ["P1", "Capital Velocity (ROE, ROIC, turnover)", "D1 TrueData Corporate + D3 FMP + D5 SimFin"],
        ["P2", "Altman Z + Beneish M", "D1, D3, D5 — computed in Laravel from statements"],
        ["P3", "Sloan Ratio", "D1, D3, D5 — cash flow vs net income"],
        ["P4", "DSO + promoter pledge", "D1, D3 + E1 TrueData SHP"],
        ["P5", "CAGR + PEG + FCFF/FCFE", "D1, D3, D5 + A/B price data"],
        ["P6", "VPVR / HVN support", "A1, A3, B1 — OHLCV volume history"],
        ["P7", "Mansfield RS vs index", "C1 Nifty 50 + C2 S&P 500"],
        ["P8", "VCP — ATR/Close percentile", "A1, A3, B1 — daily/minute bars"],
        ["P9", "Delivery % / block trades", "A2/A3 DelVolume + B2 Polygon trades"],
        ["P10", "ADX trend maturity", "A1, B1 — computed in Laravel from OHLCV"],
    ])

    doc.add_heading("5. Consolidated Budget — MVP Scenario", level=1)
    doc.add_paragraph("Pilot: ~50 India + ~50 US symbols | Internal use | No client redistribution")
    add_table(doc, ["Category", "APIs Included", "Monthly Low (₹)", "Monthly High (₹)", "One-Time (₹)"], [
        ["A — India market", "A1–A4", "10,000", "25,000", "50,000 – 1,50,000"],
        ["B — US market", "B1–B3", "16,716", "16,716", "0"],
        ["C — Indices", "C1–C2", "0", "0", "0"],
        ["D — Fundamentals", "D1, D3, D5, D6", "12,956", "25,920", "0"],
        ["E — India pledge", "E1 (+ E2 if needed)", "0", "5,000", "0"],
        ["F — News / RAG", "F1–F4", "4,200", "4,200", "0"],
        ["G — Gemini AI", "G1–G2", "2,000", "7,000", "0"],
        ["TOTAL", "—", "45,872", "83,836", "50,000 – 1,50,000"],
    ])
    add_table(doc, ["MVP Year 1 Summary", "Low (₹)", "High (₹)"], [
        ["Monthly recurring × 12", "5,50,464", "10,06,032"],
        ["One-time backfill", "50,000", "1,50,000"],
        ["Subtotal (before GST)", "6,00,464", "11,56,032"],
        ["GST @ 18% (est.)", "1,08,084", "2,08,086"],
        ["GRAND TOTAL YEAR 1", "7,08,548 (~₹7.1 Lakh)", "13,64,118 (~₹13.6 Lakh)"],
    ])

    doc.add_heading("6. Consolidated Budget — Production Scenario", level=1)
    doc.add_paragraph("Scale: ~200 symbols per market | Client-facing dashboard | Redistribution license likely required")
    add_table(doc, ["Category", "APIs Included", "Monthly Low (₹)", "Monthly High (₹)", "One-Time (₹)"], [
        ["A — India market", "A1–A5", "35,000", "80,000", "1,50,000 – 4,00,000"],
        ["B — US market", "B1–B4", "16,716", "25,000", "0"],
        ["C — Indices", "C1–C2", "0", "0", "0"],
        ["D — Fundamentals", "D1–D5", "35,000", "55,000", "0"],
        ["E — India pledge", "E1–E2", "0", "8,000", "0"],
        ["F — News / RAG", "F1–F5", "4,200", "41,916", "0"],
        ["G — Gemini AI", "G1–G2", "5,000", "15,000", "0"],
        ["TOTAL", "—", "95,916", "2,25,000", "1,50,000 – 4,00,000"],
    ])
    add_table(doc, ["Production Year 1 Summary", "Low (₹)", "High (₹)"], [
        ["Monthly recurring × 12", "11,50,992", "27,00,000"],
        ["One-time backfill", "1,50,000", "4,00,000"],
        ["Subtotal (before GST)", "13,00,992", "31,00,000"],
        ["GST @ 18% (est.)", "2,34,179", "5,58,000"],
        ["GRAND TOTAL YEAR 1", "15,35,171 (~₹15.4 Lakh)", "36,58,000 (~₹36.6 Lakh)"],
    ])

    doc.add_heading("7. Recommended Approval Path", level=1)
    steps = [
        "Approve MVP Low budget (~₹7.1 Lakh Year 1 incl. GST) for Phase 1 pilot.",
        "Procure TrueData trial (10 days free) + Polygon Developer trial before full commit.",
        "Confirm 20Y backfill depth and delivery % coverage with TrueData (historicaldata@truedata.in).",
        "Confirm redistribution requirements if dashboard will be client-facing (Category A5).",
        "Re-submit Production budget for approval after 90-day MVP validation.",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}")

    doc.add_heading("8. Exclusions (Not in This Budget)", level=1)
    exclusions = [
        "Category H — Email/SMS alert APIs (Resend, Twilio, MSG91).",
        "Category J — Broker/trading APIs (Kite, IBKR, Alpaca Trading).",
        "Infrastructure — AWS/R2, VPS hosting, Sentry (optional, separate infra budget).",
        "Development team cost, Laravel hosting, domain, SSL.",
        "US market data display license beyond Polygon non-professional tier.",
    ]
    for e in exclusions:
        doc.add_paragraph(e, style="List Bullet")

    doc.add_heading("9. Client Approval", level=1)
    add_table(doc, ["Approval Item", "Approved (Yes/No)", "Approved Budget (₹)", "Sign / Date"], [
        ["MVP Scenario — Year 1 API budget", "", "", ""],
        ["Production Scenario — Year 1 API budget (deferred)", "", "", ""],
        ["TrueData as India data vendor", "", "", ""],
        ["Polygon.io as US data vendor", "", "", ""],
        ["Gemini as AI/RAG provider", "", "", ""],
    ])
    doc.add_paragraph()
    doc.add_paragraph("Authorised Signatory: _________________________    Date: _______________")
    doc.add_paragraph("Client Name / Organisation: _________________________")

    doc.add_paragraph()
    footer = doc.add_paragraph("— End of Document —")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


if __name__ == "__main__":
    doc = build()
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")
