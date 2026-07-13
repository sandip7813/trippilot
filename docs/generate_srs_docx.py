#!/usr/bin/env python3
"""Generate Hybrid Quant Engine SRS & Project Plan DOCX."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "Hybrid-Quant-Engine-SRS.docx")


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_meta_table(doc: Document) -> None:
    rows = [
        ("Document Version", "1.1"),
        ("Date", "13 July 2026"),
        ("Project", "Hybrid Quant Engine (Analysis, Backtest & Signals)"),
        ("Markets", "India (NSE/BSE), US (NYSE/NASDAQ)"),
        ("Trading Style", "Buy-side research only — no intraday; no order execution"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    doc.add_paragraph()


def add_bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered_list(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.rows[i + 1].cells[j].text = val
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    set_doc_defaults(doc)

    add_title(doc, "HYBRID QUANT ENGINE")
    p = doc.add_paragraph("Software Requirements Specification (SRS) & Project Plan")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    doc.add_paragraph()

    add_meta_table(doc)

    # 1 Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "Build an automated buy-side quantitative research and signal platform that evaluates "
        "equities on a 10-Point Confluence Matrix (fundamental + technical + order-flow), runs "
        "20-year backtests, streams live market data in two operational shifts, and publishes "
        "actionable signals and alerts. The portal does not place, route, modify, or cancel "
        "any buy or sell orders — all execution remains outside this system."
    )

    # 2 Objectives
    doc.add_heading("2. Objectives", level=1)
    add_numbered_list(
        doc,
        [
            "Ingest and maintain 20 years (2006–2026) of corporate-action-adjusted market and fundamental data.",
            "Score every ticker on a deterministic 0–10 binary confluence scale.",
            "Backtest signal rules with 0.25% round-trip friction/slippage (simulated).",
            "Stream live data and recalculate scores in near real time (bar-level).",
            "Generate entry, hold, and exit signals with suggested risk parameters (informational only).",
            "Provide an operator dashboard for monitoring, research (RAG), watchlists, and alerts.",
        ],
    )

    # 3 Scope
    doc.add_heading("3. Scope", level=1)
    doc.add_heading("3.1 In Scope", level=2)
    add_bullet_list(
        doc,
        [
            "Historical and live market data ingestion (WebSocket feeds).",
            "10-point scoring engine (P1–P10).",
            "Backtest engine with friction model (simulated trades).",
            "Virtual / paper portfolio for signal tracking (no broker connection).",
            "Signal alerts: Strong Buy, Hold, Exit Warning.",
            "Suggested stop-loss and target levels (display only; not sent to brokers).",
            "Simulated risk metrics and circuit-breaker warnings.",
            "Admin dashboard (Blade, jQuery, Bootstrap).",
            "AI research assistant (Gemini + RAG over filings and news).",
            "Audit logs and signal lifecycle history.",
        ],
    )
    doc.add_heading("3.2 Out of Scope", level=2)
    add_bullet_list(
        doc,
        [
            "Buy or sell order placement, modification, or cancellation.",
            "Broker API integration (Kite, IBKR, Alpaca, etc.).",
            "OCO / bracket order execution.",
            "Intraday / scalping strategies.",
            "Short selling / sell-side operations.",
            "Mobile native apps (Phase 1).",
            "Regulatory registration on behalf of the client.",
        ],
    )

    # 4 Technology Stack
    doc.add_heading("4. Technology Stack", level=1)
    add_table(
        doc,
        ["Layer", "Technology", "Purpose"],
        [
            ["Backend", "Laravel 13", "Auth, queues, scheduling, scoring, signals, APIs"],
            ["Queue / Workers", "Laravel Horizon, Supervisor", "Backtests, EOD refresh, stream commands"],
            ["Cache / Pub-Sub", "Redis", "Live scores, alert state, pub/sub"],
            ["Transactional DB", "MySQL", "Users, watchlists, virtual portfolios, signals, audit"],
            ["Time-Series DB", "MongoDB", "OHLCV, fundamentals, indicators, score history"],
            ["Frontend", "Blade, jQuery, Bootstrap 5", "Dashboard, scanners, reports, alerts"],
            ["AI — Primary", "Google Gemini (Flash/Pro)", "Summarization, Q&A, research chat"],
            ["AI — RAG", "LLM + vector store", "Filings, news, earnings transcript search"],
            ["AI — Optional", "OpenAI GPT-4o / Claude", "Structured extraction from complex filings (Phase 2)"],
        ],
    )
    doc.add_paragraph(
        "Recommendation: Use Gemini as the default LLM for cost-effective research chat. "
        "Keep all scoring and signal rules deterministic in PHP. Use RAG for document-backed "
        "answers with citations. AI must never trigger trades or override the scoring engine."
    )

    # 5 Architecture
    doc.add_heading("5. System Architecture", level=1)
    doc.add_paragraph(
        "Blade + jQuery + Bootstrap Dashboard\n"
        "        │\n"
        "Laravel Application (Auth, Scoring, Signals, Backtest, RAG, Alerts)\n"
        "        ├── MySQL (users, watchlists, virtual portfolios, signals, audit)\n"
        "        ├── MongoDB (bars, fundamentals, indicators, score history)\n"
        "        └── Redis (live scores, alert flags, pub/sub)\n"
        "        │\n"
        "Long-Running Artisan Workers (Supervisor)\n"
        "  • markets:stream-india (IST 09:15–15:30)\n"
        "  • markets:stream-us (IST 19:00–02:30)\n"
        "  • refresh:fundamentals (post-market)\n"
        "        │\n"
        "External: Market Data APIs only (no broker APIs)"
    )

    # 6 Functional Requirements
    doc.add_heading("6. Functional Requirements", level=1)

    doc.add_heading("6.1 Data & Backtest Engine (FR-DATA)", level=2)
    add_table(
        doc,
        ["ID", "Requirement"],
        [
            ["FR-DATA-01", "Store 20-year daily OHLCV (2006–2026), split/bonus/dividend adjusted"],
            ["FR-DATA-02", "Ingest India daily delivery %; US block trade metrics"],
            ["FR-DATA-03", "Stream live data via WebSocket during defined market shifts"],
            ["FR-DATA-04", "Recalculate technical indicators in memory on each new bar"],
            ["FR-DATA-05", "Refresh fundamentals and sector data at post-market intervals"],
            ["FR-DATA-06", "Backtest with 0.25% round-trip friction on simulated trades"],
        ],
    )

    doc.add_heading("6.2 10-Point Confluence Scoring (FR-SCORE)", level=2)
    doc.add_paragraph("Each point is binary: 1 = Pass, 0 = Fail. Total score = sum (0–10).")
    add_table(
        doc,
        ["ID", "Point", "Pass Criteria"],
        [
            ["P1", "Capital Velocity", "5Y avg ROE > 15% AND ROIC > 15%, expanding asset turnover"],
            ["P2", "Insolvency & Hygiene", "Altman Z > 3.0 AND Beneish M-Score < −1.60"],
            ["P3", "Earnings Cash Quality", "Sloan Ratio between −10% and +10%"],
            ["P4", "Operational Continuity", "DSO stable/decreasing YoY; India: promoter pledge < 10% (reject if > 15%)"],
            ["P5", "Growth & Compounding", "≥ 15% 3Y and 5Y CAGR; PEG < 1.2; FCFF > FCFE"],
            ["P6", "Structural Support", "Price above and holding Volume Profile HVN (VPVR)"],
            ["P7", "Alpha Acceleration", "Mansfield RS > 0 vs Nifty 50 / S&P 500; 3-week slope positive"],
            ["P8", "Volatility Pocket (VCP)", "ATR/Close in lowest 15th percentile of 6-month range"],
            ["P9", "Order Flow", "Breakout vol > 1.5× 20-day SMA; India delivery > 55% OR US expanding block ratio"],
            ["P10", "Trend Maturity", "ADX 20–40, +DI > −DI; reject if ADX > 50"],
        ],
    )
    doc.add_paragraph("FR-SCORE-11: Scoring must be 100% deterministic — identical inputs yield identical outputs in backtest and live.")

    doc.add_heading("6.3 Signal Lifecycle — No Execution (FR-SIGNAL)", level=2)
    doc.add_paragraph(
        "The system publishes signals and alerts only. No orders are sent to any broker or exchange."
    )
    add_table(
        doc,
        ["Score", "Signal", "System Action"],
        [
            ["≥ 8", "STRONG BUY", "Alert + dashboard flag; show suggested position size and SL/TP levels (informational)"],
            ["4–7", "HOLD", "Maintain watchlist status; no alert unless score changes"],
            ["≤ 3", "EXIT WARNING", "High-priority alert; flag structural breakdown (user acts externally)"],
        ],
    )
    doc.add_paragraph(
        "Suggested sizing (virtual portfolio only): volatility-adjusted fixed fractional risk; "
        "max 20% per symbol, 30% per sector — displayed as guidance, not executed."
    )
    doc.add_paragraph(
        "Suggested brackets (display only): minimum 1:2 or 1:3 risk-reward using chart resistance or ATR projections."
    )

    doc.add_heading("6.4 Simulated Risk & Circuit Breakers (FR-RISK)", level=2)
    doc.add_paragraph("Applied to virtual/paper portfolios and new-signal gating — not live orders.")
    add_table(
        doc,
        ["ID", "Rule"],
        [
            ["FR-RISK-01", "Simulated open risk (sum of 1R) ≤ 6% of virtual capital; suppress new Strong Buy alerts at limit"],
            ["FR-RISK-02", "At +1R move, update virtual stop to breakeven in paper portfolio"],
            ["FR-RISK-03", "Virtual equity −5% in one calendar month → halt new Strong Buy alerts for remainder of month"],
            ["FR-RISK-04", "Manual alert mute / kill-switch on dashboard"],
        ],
    )

    doc.add_heading("6.5 AI / RAG Layer (FR-AI)", level=2)
    add_table(
        doc,
        ["ID", "Requirement"],
        [
            ["FR-AI-01", "RAG over earnings filings, annual reports, news (India + US)"],
            ["FR-AI-02", "Natural-language Q&A: e.g. “Why did symbol X score 7/10?”"],
            ["FR-AI-03", "AI must not compute scores, size positions, or trigger signals"],
            ["FR-AI-04", "All AI responses cite source document chunks"],
            ["FR-AI-05", "Optional: flag anomalies for human review (e.g. pledge increase)"],
        ],
    )

    doc.add_heading("6.6 Dashboard (FR-UI)", level=2)
    add_bullet_list(
        doc,
        [
            "Universe scanner with P1–P10 score breakdown",
            "Watchlists and virtual portfolio tracker",
            "Signal feed (Strong Buy / Hold / Exit Warning)",
            "Suggested SL/TP levels (read-only)",
            "Backtest run launcher and results viewer",
            "Simulated risk meter and circuit-breaker status",
            "RAG research panel",
            "Market shift status (India / US)",
        ],
    )

    # 7 NFR
    doc.add_heading("7. Non-Functional Requirements", level=1)
    add_table(
        doc,
        ["ID", "Category", "Requirement"],
        [
            ["NFR-01", "Performance", "Score update within ≤ 5 seconds of new bar"],
            ["NFR-02", "Availability", "99.5% uptime during market shifts"],
            ["NFR-03", "Security", "RBAC, encrypted credentials for data vendors, full audit trail"],
            ["NFR-04", "Reproducibility", "Backtest results reproducible from stored inputs + versioned logic"],
            ["NFR-05", "Scalability", "500+ symbols per market (Phase 2)"],
            ["NFR-06", "Compliance", "Immutable signal and score event log"],
        ],
    )

    # 8 Data Model
    doc.add_heading("8. Data Model (Summary)", level=1)
    doc.add_paragraph("MySQL: users, watchlists, virtual_portfolios, virtual_positions, signals, alert_rules, "
                      "risk_limits, circuit_breaker_events, backtest_runs, audit_logs")
    doc.add_paragraph("MongoDB: daily_bars, delivery_data, block_trades, fundamental_snapshots, "
                      "indicator_series, score_history, rag_document_chunks")
    doc.add_paragraph("Redis: live_scores:{symbol}, virtual_portfolio_risk, monthly_halt_flag")

    # 9 Shifts
    doc.add_heading("9. Operational Shifts", level=1)
    add_table(
        doc,
        ["Shift", "Window (IST)", "Market"],
        [
            ["1", "09:15 – 15:30", "India (NSE/BSE)"],
            ["2", "19:00 – 02:30", "US (NYSE/NASDAQ)"],
        ],
    )
    doc.add_paragraph("Post-market: fundamentals refresh, sector recalc, backtest jobs.")

    # 10 Project Plan
    doc.add_heading("10. Project Plan (Phases)", level=1)
    add_table(
        doc,
        ["Phase", "Duration", "Deliverables"],
        [
            ["P0 — Foundation", "3 weeks", "Laravel setup; MySQL/MongoDB/Redis; auth; instrument master; data vendor (1 market)"],
            ["P1 — Scoring & Backtest", "4 weeks", "P1–P10 services; 20Y EOD ingest; backtest engine; Pest tests"],
            ["P2 — Dashboard", "3 weeks", "Blade/jQuery/Bootstrap: scanner, score detail, backtest reports"],
            ["P3 — Virtual Portfolio", "2 weeks", "Paper tracking; simulated risk rules; signal audit log"],
            ["P4 — Live Streaming", "3 weeks", "WebSocket workers; Redis live scores; shift scheduling"],
            ["P5 — Alerts & Signals", "2 weeks", "Strong Buy / Exit alerts; email or in-app notifications"],
            ["P6 — AI / RAG", "3 weeks", "Document pipeline; Gemini chat; cited Q&A"],
            ["P7 — Hardening & UAT", "2 weeks", "Load test, reconciliation, operator training"],
        ],
    )
    doc.add_paragraph("Total estimated duration: ~22 weeks (~5.5 months)")

    # 11 Milestones
    doc.add_heading("11. Milestones & Acceptance Criteria", level=1)
    add_table(
        doc,
        ["Milestone", "Acceptance"],
        [
            ["M1 — Data ready", "20Y adjusted bars for pilot universe (≥ 50 symbols)"],
            ["M2 — Scoring validated", "P1–P10 Pest tests pass; spot-check vs spreadsheet"],
            ["M3 — Backtest complete", "Walk-forward run with friction; exportable simulated trade log"],
            ["M4 — Virtual portfolio", "30-day signal tracking with alerts; zero broker calls"],
            ["M5 — Live signals", "Real-time score updates and alerts during market shifts"],
            ["M6 — RAG live", "Filing-backed Q&A with citations"],
        ],
    )

    # 12 Assumptions
    doc.add_heading("12. Assumptions & Dependencies", level=1)
    add_bullet_list(
        doc,
        [
            "Client procures licensed market data (EOD, delivery, block trades, fundamentals).",
            "No broker accounts or API keys are connected to this portal.",
            "Users execute trades manually on external platforms.",
            "Scoring uses daily (or 5m/15m) bars — not tick-level HFT.",
            "Initial pilot universe: 50–100 liquid symbols per market.",
        ],
    )

    # 13 Risks
    doc.add_heading("13. Risks", level=1)
    add_table(
        doc,
        ["Risk", "Mitigation"],
        [
            ["Data quality / corporate actions", "Adjustment pipeline + reconciliation reports"],
            ["PHP backtest speed", "Horizon parallel jobs; overnight runs"],
            ["AI hallucination", "RAG with citations; AI isolated from scoring/signals"],
            ["Overfitting", "Walk-forward validation; out-of-sample holdout"],
            ["User expects auto-trading", "Explicit UI disclaimer; no broker integration by design"],
        ],
    )

    # 14 Team
    doc.add_heading("14. Suggested Team", level=1)
    add_table(
        doc,
        ["Role", "Count"],
        [
            ["Laravel backend developer", "1–2"],
            ["Frontend (Blade/jQuery)", "1"],
            ["Quant / domain advisor", "1 (part-time)"],
            ["DevOps", "1 (part-time)"],
        ],
    )

    doc.add_paragraph()
    footer = doc.add_paragraph("— End of Document —")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


if __name__ == "__main__":
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")
